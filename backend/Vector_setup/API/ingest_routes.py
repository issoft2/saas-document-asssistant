from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status
from pydantic import BaseModel
from typing import List, Optional
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
import uuid
import pandas as pd


from Vector_setup.base.db_setup_management import (
    MultiTenantChromaStoreManager,
    TenantCollectionConfigRequest,
    CollectionCreateRequest,
)
from Vector_setup.user.auth_jwt import get_current_user
from Vector_setup.base.auth_models import UserOut

router = APIRouter()

# Single shared store instance
vector_store = MultiTenantChromaStoreManager("./chromadb_multi_tenant")


def get_store() -> MultiTenantChromaStoreManager:
    return vector_store


# ---------- Schemas for responses ----------

class TenantCollectionConfigOut(BaseModel):
    status: str
    tenant_id: str
    collection_name: str


class CompanyOut(BaseModel):
    tenant_id: str
    display_name: str | None = None


class CollectionOut(BaseModel):
    tenant_id: str
    collection_name: str
    doc_count: int


# ---------- Role helpers ----------

def require_vendor(current_user: UserOut = Depends(get_current_user)) -> UserOut:
    if current_user.role != "vendor":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only vendor can perform this action.",
        )
    return current_user


ALLOWED_COLLECTION_CREATORS = {"hr", "executive"}


def require_collection_creator(
    current_user: UserOut = Depends(get_current_user),
) -> UserOut:
    """
    Users allowed to create collections: hr, executive.
    Vendor is explicitly blocked here (even though vendor is 'super' elsewhere).
    """
    if current_user.role in ALLOWED_COLLECTION_CREATORS:
        return current_user

    if current_user.role == "vendor":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Vendor cannot create collections for companies.",
        )

    raise HTTPException(
        status_code=status.HTTP_403_FORBIDDEN,
        detail="Only HR or Executive can create collections.",
    )


def require_uploader(
    current_user: UserOut = Depends(get_current_user),
) -> UserOut:
    """
    Users allowed to upload documents: hr, executive.
    Extend if you want management etc.
    """
    if current_user.role in ALLOWED_COLLECTION_CREATORS:
        return current_user

    raise HTTPException(
        status_code=status.HTTP_403_FORBIDDEN,
        detail="Only HR or Executive can upload documents.",
    )


# ---------- Admin configuration APIs ----------

@router.post("/companies/configure", response_model=TenantCollectionConfigOut)
def configure_company_and_collection(
    req: TenantCollectionConfigRequest,
    store: MultiTenantChromaStoreManager = Depends(get_store),
    current_user: UserOut = Depends(require_vendor),
):
    """
    Configure a company and its first collection in a single call.
    Only vendor can create/provision companies.
    """
    # At this point require_vendor has already ensured role == "vendor".
    result = store.configure_tenant_and_collection(req)
    return TenantCollectionConfigOut(
        status=result["status"],
        tenant_id=result["tenant_id"],
        collection_name=result["collection_name"],
    )


# ---- Extra schema for creation ----

class CollectionCreateIn(BaseModel):
    name: str  # just the collection/policy name


@router.post("/collections", response_model=CollectionOut)
def create_collection_for_current_tenant(
    req: CollectionCreateIn,
    store: MultiTenantChromaStoreManager = Depends(get_store),
    current_user: UserOut = Depends(require_collection_creator),
):
    """
    HR/Executive creates a collection for their own tenant.
    Tenant is taken from the authenticated user, not from the client.
    Vendor cannot create collections.
    """
    tenant_id = current_user.tenant_id

    result = store.create_collection(
        CollectionCreateRequest(
            tenant_id=tenant_id,
            collection_name=req.name,
        )
    )
    return CollectionOut(
        tenant_id=tenant_id,
        collection_name=result["collection_name"],
        doc_count=result["document_count"],
    )


@router.get("/companies", response_model=List[CompanyOut])
def list_companies(
    store: MultiTenantChromaStoreManager = Depends(get_store),
    current_user: UserOut = Depends(get_current_user),
):
    """
    List all companies/tenants that currently have collections.
    - Vendor: sees all companies.
    - Other users: see only their own company.
    """
    all_tenants = store.list_companies()  # list[{ "tenant_id", "display_name"}]
    my_tenant_id = current_user.tenant_id

    if current_user.role == "vendor":
        return [CompanyOut(**t) for t in all_tenants]

    return [
        CompanyOut(**t)
        for t in all_tenants
        if t["tenant_id"] == my_tenant_id
    ]


@router.get("/companies/{tenant_id}/collections", response_model=List[CollectionOut])
def list_company_collections(
    tenant_id: str,
    store: MultiTenantChromaStoreManager = Depends(get_store),
    current_user: UserOut = Depends(get_current_user),
):
    """
    List all collections for a given tenant (names without tenant prefix).
    - Vendor: can see collections for any tenant.
    - Other users: only allowed to view their own tenant.
    """
    if current_user.role != "vendor" and tenant_id != current_user.tenant_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not allowed to access this tenant",
        )

    names = store.list_collections(tenant_id)
    # If you don't have doc counts for now, set to 0.
    return [
        CollectionOut(
            tenant_id=tenant_id,
            collection_name=n,
            doc_count=0,
        )
        for n in names
    ]


# ---------- Document upload (production) ----------

@router.post("/documents/upload")
async def upload_document(
    tenant_id: str = Form(...),
    collection_name: str = Form(...),
    title: Optional[str] = Form(None),
    doc_id: Optional[str] = Form(None),
    file: UploadFile = File(...),
    store: MultiTenantChromaStoreManager = Depends(get_store),
    current_user: UserOut = Depends(require_uploader),
):
    """
    Upload a document file and index it into the tenant's collection.

    Rules:
    - Only HR/Executive can upload (require_uploader).
    - Must upload only into their own tenant.
    """
    # Enforce tenant isolation
    if tenant_id != current_user.tenant_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not allowed to upload into this tenant.",
        )

    # Basic validation
    if not tenant_id.replace("-", "").replace("_", "").isalnum():
        raise HTTPException(status_code=400, detail="Invalid tenant_id")

    if not collection_name.replace("-", "").replace("_", "").isalnum():
        raise HTTPException(status_code=400, detail="Invalid collection name")

    # Read raw bytes
    raw_bytes = await file.read()

    text = extract_text_from_upload(file.filename, raw_bytes)
    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="No text could be extracted from the document",
        )

    # Generate doc_id if not supplied
    final_doc_id = doc_id or str(uuid.uuid4())

    # Build document-level metadata (for later source display)
    metadata = {
        "filename": file.filename,
        "title": title or file.filename,
        "content_type": file.content_type,
        "size_bytes": len(raw_bytes),
        "tenant_id": tenant_id,
        "collection": collection_name,
    }

    # Delegate to vector store (chunking + embeddings)
    result = await store.add_document(
        tenant_id=tenant_id,
        collection_name=collection_name,
        doc_id=final_doc_id,
        text=text,
        metadata=metadata,
    )

    if result.get("status") != "ok":
        raise HTTPException(
            status_code=500,
            detail=result.get("message", "Indexing failed"),
        )

    return result


# ---------- Pdf Text extraction helpers ----------

def _extract_pdf_with_pymupdf(raw_bytes: bytes) -> str:
    parts: List[str] = []

    with fitz.open(stream=raw_bytes, filetype="pdf") as doc:  # type: ignore[arg-type]
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                parts.append(text.strip())

            try:
                tables = page.find_tables()
            except Exception:
                tables = None

            if tables:
                for table in tables.tables:
                    try:
                        md = table.to_markdown()
                    except Exception:
                        rows = []
                        for row in table.extract():
                            rows.append(" | ".join(cell or "" for cell in row))
                        md = "\n".join(rows)

                    if md.strip():
                        parts.append(md.strip())

    return "\n\n".join(parts)


# ---------- Excel Text extraction helpers ----------
def _extract_excel_with_pandas(raw_bytes: bytes, filename: str) -> str:
    """
    Extracts human‑readable text from an Excel file.

    - Reads all sheets using pandas.
    - Skips empty sheets.
    - For each non‑NaN cell, emits `Column: Value`.
    - Groups cells by row, separated with ` | `.
    - Separates sheets with a blank line.
    """
    buffer = BytesIO(raw_bytes)

    # sheet_name=None -> dict[str, DataFrame] for all sheets.[web:251][web:316]
    sheets = pd.read_excel(buffer, sheet_name=None, engine="openpyxl")

    sheet_chunks: List[str] = []

    for sheet_name, df in sheets.items():
        if df.empty:
            continue

        row_lines: List[str] = []

        # iterrows is fine here; Excel sheets are usually modest in size.
        for _, row in df.iterrows():
            parts: List[str] = []
            for col, val in row.items():
                if pd.isna(val):
                    continue

                parts.append(f"{col}: {val}")

            if parts:
                row_lines.append("  |  ".join(parts))

        if not row_lines:
            continue

        sheet_text = f"Sheet: {sheet_name}\n" + "\n".join(row_lines)
        sheet_chunks.append(sheet_text)

    # Join all sheets into a single string so downstream code always gets str.
    return "\n\n".join(sheet_chunks)

# ---------- Main text extraction function ----------
def extract_text_from_upload(filename: str, raw_bytes: bytes) -> str:
    name = filename.lower()

    if name.endswith(".md") or name.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        return _extract_pdf_with_pymupdf(raw_bytes)
    
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        return _extract_excel_with_pandas(raw_bytes, name)

    if name.endswith(".docx"):
        doc = Document(BytesIO(raw_bytes))
        parts: List[str] = []

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_texts = [p.text for p in cell.paragraphs if p.text]
                    cells.append(" ".join(cell_texts))
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    return ""
