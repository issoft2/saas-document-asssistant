from io import BytesIO
from docx import Document
from typing import List
import csv
from io import StringIO


from Vector_setup.services.extracting_excel_document_service import _extract_excel_with_pandas
from Vector_setup.services.extracting_pdf_document_service import _extract_pdf_with_pymupdf

import logging

logger = logging.getLogger(__name__)


# ---------- Main text extraction function ----------
def extract_text_from_upload(filename: str, raw_bytes: bytes) -> str:
    name = filename.lower()

    if name.endswith(".md") or name.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        return _extract_pdf_with_pymupdf(raw_bytes)
    
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        return _extract_excel_with_pandas(raw_bytes, name)
    
    if name.endswith(".csv"):
        # Decode bytes → text
        text = raw_bytes.decode("utf-8", errors="ignore")
        buffer = StringIO(text)

        reader = csv.reader(buffer)
        rows = []
        for row in reader:
            # join columns with separator; tweak as you like
            rows.append(" | ".join(col.strip() for col in row if col is not None))

        return "\n".join(rows) 

    if name.endswith(".docx"):
        doc = Document(BytesIO(raw_bytes))
        parts: List[str] = []

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_texts = [p.text for p in cell.paragraphs if p.text]
                    cells.append(" ".join(cell_texts))
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    
    logger.warning("Unsupported file type for text extraction: %s", filename)
    return ""